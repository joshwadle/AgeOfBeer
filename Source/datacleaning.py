import pandas as pd
from docx import Document
from os import listdir, path
from comment_nlp_py3 import nlpComments
import re
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer



class Data_Loading(object):
    def init(self):
        if not 'nlp' in locals():
            print("Loading English Module...")
            nlp = spacy.load('en')

    def Document_List(self):
        '''
        Creates a list of "Documents" that can be passed into Create_Dataframe
        It also passes out a shorter list to test functions on.
        '''
        AVLlst = [Document('SLS/AVL Brands/FT AVL SLS 1 MO AUG 25 16.docx'),
        Document('SLS/AVL Brands/FT AVL SLS 1 MO JUNE 6 16.docx'),
        Document('SLS/AVL Brands/FT AVL SLS 2 mo SEPT 26 16.docx'),
        Document('SLS/AVL Brands/FT AVL SLS 3 MO OCT 24 16.docx'),
        Document('SLS/AVL Brands/FT AVL SLS 4 MO DEC 1 16.docx')]
        Document('SLS/AVL Brands/FT AVL SLS 1 MO AUG 25 16.docx'),
        Document('SLS/AVL Brands/FT AVL SLS 1 MO JUNE 6 16.docx'),
        Document('SLS/AVL Brands/FT AVL SLS 2 mo SEPT 26 16.docx'),
        Document('SLS/AVL Brands/FT AVL SLS 3 MO OCT 24 16.docx'),
        Document('SLS/AVL Brands/FT AVL SLS 4 MO DEC 1 16.docx')

        Ciderlst = [Document('SLS/Cider/{}'.format(name)) for name in listdir('SLS/Cider')]
        GRlst = [Document('SLS/Gluten reduced/{}'.format(name)) for name in listdir('SLS/Gluten reduced')]
        Hoplst = [Document('SLS/Hop Kitchen/{}'.format(name)) for name in listdir('SLS/Hop Kitchen')]
        LoFlst = [Document('SLS/LoF/{}'.format(name)) for name in listdir('SLS/LoF')]
        Pilotlst = [Document('SLS/Pilot Brewery/{}'.format(name)) for name in listdir('SLS/Pilot Brewery')]
        Reglst = [Document('SLS/Regular Brands/{}'.format(name)) for name in listdir('SLS/Regular Brands')]
        Seslst = [Document('SLS/Seasonal/{}'.format(name)) for name in listdir('SLS/Seasonal')]

        self.doclst =  Ciderlst+GRlst+Hoplst+LoFlst+Pilotlst+Seslst+Reglst+AVLlst
        return self.doclst

    def Create_Dataframe(self):
        '''
        Takes a list of Documents and creates a Pandas DataFrame
        '''
        Beerlst =[]
        RegCodelst = []
        Namelst = []
        Samplst = []
        Comments = []
        AgeName = []
        Questionlst = []
        total = 0
        for doc in self.doclst:
            total+=1
            print(doc, total)
            for i in range(len(doc.tables)):
                if re.findall(r"(\d*\.*\d+)\s*[MO|mo]", doc.paragraphs[0].text):
                    age = float(re.findall(r"(\d*\.*\d+)\s*[MO|mo]", doc.paragraphs[0].text)[0])
                else:
                    age = 0
                if len(doc.tables)<=15 or i < 13:
                    if i%3==2:
                        if len(doc.tables[i-2].rows) >= 3:
                            if len(doc.tables[i-2].rows[2].cells)>=2:
                                question = doc.tables[i-2].cell(2,1).text
                        # print(i)
                        for row in doc.tables[i].rows:
                            if len(row.cells)==6:
                                if row.cells[0].text != 'RegCode' and row.cells[0].text != '':
                                    Questionlst.append(question)
                                    AgeName.append(age)
                                    Beerlst.append(doc.tables[1].cell(1,1).text)
                                    RegCodelst.append(row.cells[0].text)
                                    Namelst.append(row.cells[1].text)
                                    Samplst.append(row.cells[4].text)
                                    Comments.append(row.cells[5].text)
                elif i > 15 and i%3==0:
                    if len(doc.tables[i-2].rows) >= 3:
                        if len(doc.tables[i-2].rows[2].cells)>=2:
                            question = doc.tables[i-2].cell(2,1).text
                    # print(i)
                    for row in doc.tables[i].rows:
                        if len(row.cells)==6:
                            if row.cells[0].text != 'RegCode'and row.cells[0].text != '':
                                Questionlst.append(question)
                                AgeName.append(age)
                                Beerlst.append(doc.tables[1].cell(1,1).text)
                                RegCodelst.append(row.cells[0].text)
                                Namelst.append(row.cells[1].text)
                                Samplst.append(row.cells[4].text)
                                Comments.append(row.cells[5].text)

        lst_of_lst = [Beerlst, RegCodelst, Namelst, Samplst, Comments, AgeName, Questionlst]
        df = pd.DataFrame(lst_of_lst)
        df = df.transpose()
        df.columns = ['Beer', 'RegCode', 'Name', 'Sample', 'Comments', 'Age', 'Question']
        self.df = df
        return df

    def Clean_Question(self):
        '''
        This changes the Question column and cleans it
        '''
        QuestionDict = {}
        QuestionDict['Clarity Comment  '] = 'clarity'
        QuestionDict['Visual Comment  '] ='clarity'
        QuestionDict['Aroma Comment  '] = 'aroma'
        QuestionDict['Aroma Comment   '] = 'aroma'
        QuestionDict['Aroma comment  '] = 'aroma'
        QuestionDict['Mouthfeel/Body Comment  '] = 'body'
        QuestionDict['Mouthfeel Comments  '] = 'body'
        QuestionDict['Taste Comment  '] = 'taste'
        QuestionDict['Flavor Comment  '] = 'taste'
        QuestionDict['Final Comments  '] ='final'
        QuestionDict['Not Sellable Comment  '] ='final'
        QuestionDict['Sellable Comment  '] ='final'
        QuestionDict['Wild Wild Dubbel Keg'] ='final'
        QuestionDict['fresh ttb  '] ='final'
        QuestionDict['True to Target  '] ='final'
        QuestionDict['not true to target  '] ='final'
        QuestionDict['Snap shot'] ='final'
        new_question=[]

        for title in self.df.Question:
            new_question.append(QuestionDict[title])
        self.df['Question'] = new_question
        return new_question

    def Trustworthiness(self):
        '''
        This puts in the Trustworthiness category in the DataFrame. This must be used after the
        Clean_Names and Clean_Question functions. otherwise they will all be weighted evenly.
        '''
        self.trust = pd.read_pickle('trustworthiness_ratings.pkl')
        check = set(trust.index)
        new_trust = []
        for row in range(self.df.shape[0]):
            Name = self.df.Name[row]
            Question = self.df.Question[row]
            if Name in check:
                if Question == 'clarity':
                    new_trust.append(trust['clarity'][Name].mean())
                    # print(trust['clarity'][Name].mean())
                elif Question == 'aroma':
                    new_trust.append(trust['aroma'][Name].mean())
                    # print(trust['aroma'][Name].mean())
                elif Question == 'taste':
                    new_trust.append(trust['taste'][Name].mean())
                    # print(trust['taste'][Name].mean())
                elif Question == 'body':
                    new_trust.append(trust['body'][Name].mean())
                    # print(trust['body'][Name].mean())
                else:
                    new_trust.append(trust['mean'][Name].mean())
                    # print(trust['mean'][Name].mean())
            else:
                new_trust.append(2.3857758620689653)
        df['Trust'] = new_trust
        return new_trust

    def Clean_Names(self):
        '''
        This cleans up the names in the dataframe so that they are uniform
        '''
        name_dict = pd.read_pickle('name_clean_dict.pkl')
        check = set(name_dict.keys())
        new_names = []
        for name in self.df.Name:
            if name not in check:
                name_dict[name.lower()]=name.lower()
            new_names.append(name_dict[name.lower()])
        self.df['Name'] = new_names
        return new_names

    def Tokenize_Clean_Strings(self):
        '''
        This will use the nlpComments class to clean the text and then it will create the tfidf friendly comment
        '''
        com = nlpComments(nlp)
        com.Fit_transform(df.Comments)
        com.Tag_comment_list()
        self.df['Comments'] = com.Make_tfidf_friendly()
        return new_comments

    def Clean_Beer_Names(self):
        '''
        This cleans the Beer names so that they are all uniform
        '''
        beer_dict = pd.read_pickle('Beer_dict.pkl')
            new_beer=[]
        for beer in df.Beer:
            new_beer.append(Beer_names[beer])
        self.df['Beer'] = new_beer
        return new_beer



if __name__ == '__main__':
    DL = Data_Loading()
    if not path.exists('data.pkl'):
        doclst = DL.Document_List()
        df = DL.Create_Dataframe(doclst)
        new_comments = DL.Tokenize_Clean_Strings(df)
        DL.Clean_Beer_Names(df)
        DL.Trustworthiness(df)
        TFIDF = DL.TfidfVectorizer()
        vectorized = TFIDF.fit_transform(new_comments)
        df['Vectorized'] = vectorized
        df.to_pickle('data.pkl')
    else:
        df = pd.read_pickle('data.pkl')
